"""
Document metadata extraction (PDF, Office documents)
"""

import os
import logging
from typing import Dict, Any, Optional

from ..config.constants import SUPPORTED_EXTENSIONS
from ..config.dependencies import OPTIONAL_DEPENDENCIES
from .base import BaseExtractor


class DocumentExtractor(BaseExtractor):
    """Extract metadata from document files (PDF, Office documents)."""
    
    @classmethod
    def can_handle(cls, file_path: str, mime_type: Optional[str] = None) -> bool:
        """Check if this extractor can handle the specified file."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in SUPPORTED_EXTENSIONS['documents']:
            return True
        
        if mime_type and any(mime in mime_type for mime in 
                            ['application/pdf', 'application/msword', 'application/vnd.openxmlformats',
                             'application/vnd.ms-excel', 'application/vnd.ms-powerpoint']):
            return True
            
        return False
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document files."""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_word(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._extract_excel(file_path)
        elif ext in ['.pptx', '.ppt']:
            return self._extract_powerpoint(file_path)
        else:
            return {'error': f"Unsupported document format: {ext}"}
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {
            'document_info': {},
            'xmp_metadata': {},
            'page_info': {},
            'security': {}
        }
        
        try:
            from PyPDF2 import PdfReader
            
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                
                # Basic document info
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if isinstance(key, str):
                            clean_key = key.lstrip('/')
                            metadata['document_info'][clean_key] = str(value)
                
                # Page information
                metadata['page_info'] = {
                    'page_count': len(reader.pages),
                    'pages': []
                }
                
                # Sample the first few pages for detailed info
                max_pages_to_sample = min(len(reader.pages), 5)
                for i in range(max_pages_to_sample):
                    page = reader.pages[i]
                    page_info = {
                        'index': i,
                        'rotation': page.get('/Rotate', 0),
                        'size': {'width': page.mediabox.width, 'height': page.mediabox.height}
                    }
                    metadata['page_info']['pages'].append(page_info)
                
                # Security information
                metadata['security'] = {
                    'encrypted': reader.is_encrypted,
                    'permissions': {
                        'printing': not reader.is_encrypted
                    }
                }
                
                if reader.is_encrypted:
                    metadata['security']['permissions'] = {
                        'printing': reader.can_print,
                        'modification': reader.can_modify,
                        'copy': reader.can_copy,
                        'annotation': reader.can_annotate,
                        'filling_forms': reader.can_fill_forms,
                        'extract_content': reader.can_extract,
                        'assemble_doc': reader.can_assemble
                    }
        except Exception as e:
            logging.error(f"Error extracting PDF metadata from {file_path}: {e}")
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_word(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word documents (DOCX, DOC)."""
        metadata = {
            'document_properties': {},
            'custom_properties': {},
            'content_stats': {}
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.docx' and OPTIONAL_DEPENDENCIES['docx']:
                import docx
                doc = docx.Document(file_path)
                
                # Core properties
                core_props = doc.core_properties
                metadata['document_properties'] = {
                    'author': core_props.author,
                    'created': str(core_props.created) if core_props.created else None,
                    'last_modified_by': core_props.last_modified_by,
                    'modified': str(core_props.modified) if core_props.modified else None,
                    'title': core_props.title,
                    'subject': core_props.subject,
                    'keywords': core_props.keywords,
                    'language': core_props.language,
                    'category': core_props.category,
                    'version': core_props.revision
                }
                
                # Content statistics
                metadata['content_stats'] = {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'page_count': 'unknown',  # Requires rendering to determine
                    'word_count': sum(len(p.text.split()) for p in doc.paragraphs),
                    'character_count': sum(len(p.text) for p in doc.paragraphs)
                }
            
            elif ext == '.doc' and OPTIONAL_DEPENDENCIES['olefile']:
                import olefile
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    
                    # Extract SummaryInformation stream
                    if ole.exists('\\x05SummaryInformation'):
                        si_stream = ole.openstream('\\x05SummaryInformation')
                        try:
                            from oletools.oleid import OleID
                            oledata = OleID(file_path)
                            indicators = oledata.check()
                            
                            for indicator in indicators:
                                if indicator.name == 'author':
                                    metadata['document_properties']['author'] = indicator.value
                                if indicator.name == 'creation_time':
                                    metadata['document_properties']['created'] = str(indicator.value)
                                if indicator.name == 'last_saved_time':
                                    metadata['document_properties']['modified'] = str(indicator.value)
                                if indicator.name == 'vba_macros':
                                    metadata['security'] = {'macros_present': indicator.value}
                        except ImportError:
                            metadata['note'] = "Limited metadata extraction for DOC (oletools not available)"
                    
                    ole.close()
            else:
                metadata['error'] = "Required dependencies not available for Word document extraction"
                
        except Exception as e:
            logging.error(f"Error extracting Word metadata from {file_path}: {e}")
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel documents (XLSX, XLS)."""
        metadata = {
            'document_properties': {},
            'custom_properties': {},
            'content_stats': {}
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.xlsx' and OPTIONAL_DEPENDENCIES['openpyxl']:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                
                # Basic properties
                metadata['document_properties'] = {
                    'creator': wb.properties.creator,
                    'created': str(wb.properties.created) if wb.properties.created else None,
                    'last_modified_by': wb.properties.lastModifiedBy,
                    'modified': str(wb.properties.modified) if wb.properties.modified else None,
                    'title': wb.properties.title,
                    'subject': wb.properties.subject,
                    'keywords': wb.properties.keywords,
                    'category': wb.properties.category
                }
                
                # Content statistics
                sheet_stats = []
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_info = {
                        'name': sheet_name,
                        'dimensions': sheet.calculate_dimension(),
                    }
                    sheet_stats.append(sheet_info)
                
                metadata['content_stats'] = {
                    'sheet_count': len(wb.sheetnames),
                    'sheets': sheet_stats
                }
                
                wb.close()
            
            elif ext == '.xls' and OPTIONAL_DEPENDENCIES['olefile']:
                # Same approach as for DOC files
                import olefile
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    
                    # Extract SummaryInformation stream
                    if ole.exists('\\x05SummaryInformation'):
                        try:
                            from oletools.oleid import OleID
                            oledata = OleID(file_path)
                            indicators = oledata.check()
                            
                            for indicator in indicators:
                                if indicator.name == 'author':
                                    metadata['document_properties']['author'] = indicator.value
                                if indicator.name == 'creation_time':
                                    metadata['document_properties']['created'] = str(indicator.value)
                                if indicator.name == 'last_saved_time':
                                    metadata['document_properties']['modified'] = str(indicator.value)
                                if indicator.name == 'vba_macros':
                                    metadata['security'] = {'macros_present': indicator.value}
                        except ImportError:
                            metadata['note'] = "Limited metadata extraction for XLS (oletools not available)"
                    
                    ole.close()
            else:
                metadata['error'] = "Required dependencies not available for Excel document extraction"
                
        except Exception as e:
            logging.error(f"Error extracting Excel metadata from {file_path}: {e}")
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_powerpoint(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PowerPoint documents (PPTX, PPT)."""
        # Basic PowerPoint extraction - could be expanded with specialized libraries
        metadata = {
            'document_properties': {},
            'content_stats': {}
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # PowerPoint files are also OLE files (for PPT) or ZIP-based (for PPTX)
            if ext == '.ppt' and OPTIONAL_DEPENDENCIES['olefile']:
                import olefile
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    
                    # Extract SummaryInformation stream (same approach as DOC/XLS)
                    if ole.exists('\\x05SummaryInformation'):
                        try:
                            from oletools.oleid import OleID
                            oledata = OleID(file_path)
                            indicators = oledata.check()
                            
                            for indicator in indicators:
                                if indicator.name == 'author':
                                    metadata['document_properties']['author'] = indicator.value
                                if indicator.name == 'creation_time':
                                    metadata['document_properties']['created'] = str(indicator.value)
                                if indicator.name == 'last_saved_time':
                                    metadata['document_properties']['modified'] = str(indicator.value)
                                if indicator.name == 'vba_macros':
                                    metadata['security'] = {'macros_present': indicator.value}
                        except ImportError:
                            metadata['note'] = "Limited metadata extraction for PPT (oletools not available)"
                    
                    ole.close()
            
            elif ext == '.pptx':
                # PPTX files are ZIP files, we can use zipfile to peek inside or python-pptx if available
                import zipfile
                with zipfile.ZipFile(file_path) as zf:
                    # Check for core properties XML file
                    if 'docProps/core.xml' in zf.namelist():
                        with zf.open('docProps/core.xml') as f:
                            import xml.etree.ElementTree as ET
                            tree = ET.parse(f)
                            root = tree.getroot()
                            
                            # Extract basic metadata from XML
                            ns = {'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
                                  'dc': 'http://purl.org/dc/elements/1.1/',
                                  'dcterms': 'http://purl.org/dc/terms/'}
                            
                            # Get creator/author
                            creator = root.find('.//dc:creator', ns)
                            if creator is not None and creator.text:
                                metadata['document_properties']['creator'] = creator.text
                            
                            # Get title
                            title = root.find('.//dc:title', ns)
                            if title is not None and title.text:
                                metadata['document_properties']['title'] = title.text
                            
                            # Get created date
                            created = root.find('.//dcterms:created', ns)
                            if created is not None and created.text:
                                metadata['document_properties']['created'] = created.text
                            
                            # Get modified date
                            modified = root.find('.//dcterms:modified', ns)
                            if modified is not None and modified.text:
                                metadata['document_properties']['modified'] = modified.text
                    
                    # Get slide count by counting XML files in slides directory
                    slide_count = len([name for name in zf.namelist() if name.startswith('ppt/slides/slide')])
                    metadata['content_stats']['slide_count'] = slide_count
            
            else:
                metadata['error'] = "Required dependencies not available for PowerPoint document extraction"
                
        except Exception as e:
            logging.error(f"Error extracting PowerPoint metadata from {file_path}: {e}")
            metadata['error'] = str(e)
        
        return metadata